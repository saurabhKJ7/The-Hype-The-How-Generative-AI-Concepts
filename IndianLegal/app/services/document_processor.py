import fitz  # PyMuPDF
from docx import Document as DocxDocument
import spacy
import uuid
from typing import List, Dict, Optional
import os
from ..models.document import Document, DocumentChunk

class DocumentProcessor:
    def __init__(self):
        self.nlp = spacy.load("en_core_web_lg")
        # Add custom legal entity patterns
        ruler = self.nlp.add_pipe("entity_ruler", before="ner")
        patterns = [
            {"label": "LEGAL_ACT", "pattern": [{"LOWER": {"IN": ["act", "section", "rule"]}}, {"LIKE_NUM": True}]},
            {"label": "LEGAL_TERM", "pattern": [{"LOWER": {"IN": ["gst", "income", "tax", "property", "court", "judgment"]}}]},
        ]
        ruler.add_patterns(patterns)

    def process_file(self, file_path: str, doc_type: str) -> Document:
        """Process a PDF or Word document and return a Document object."""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            content = self._process_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            content = self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
            
        # Create document
        doc = Document(
            id=str(uuid.uuid4()),
            title=os.path.basename(file_path),
            content=content,
            doc_type=doc_type,
            file_path=file_path
        )
        
        # Create chunks
        doc.chunks = self._create_chunks(content)
        return doc
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from Word document."""
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def _create_chunks(self, text: str, chunk_size: int = 1000) -> List[DocumentChunk]:
        """Split text into chunks and extract legal entities."""
        chunks = []
        # Split text into sentences and group into chunks
        doc = self.nlp(text)
        current_chunk = ""
        current_entities = []
        
        for sent in doc.sents:
            if len(current_chunk) + len(sent.text) > chunk_size and current_chunk:
                # Create chunk
                chunks.append(self._create_chunk(current_chunk, current_entities))
                current_chunk = ""
                current_entities = []
            
            current_chunk += sent.text + " "
            # Extract legal entities
            for ent in sent.ents:
                if ent.label_ in ["ORG", "LAW", "LEGAL_ACT", "LEGAL_TERM"]:
                    current_entities.append(ent.text)
        
        # Add final chunk if exists
        if current_chunk:
            chunks.append(self._create_chunk(current_chunk, current_entities))
        
        return chunks
    
    def _create_chunk(self, text: str, entities: List[str]) -> DocumentChunk:
        """Create a DocumentChunk with metadata."""
        return DocumentChunk(
            text=text.strip(),
            metadata={
                "length": len(text),
                "entity_count": len(entities)
            },
            legal_entities=list(set(entities))  # Remove duplicates
        ) 